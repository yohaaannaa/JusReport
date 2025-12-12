import os
import uuid
import io
import time
import traceback
from typing import Dict, Any, Optional, Tuple, List

from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, JSONResponse
from dotenv import load_dotenv

import pdfplumber
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, Inches
from pydantic import BaseModel

# pdf2image é opcional – usamos se estiver instalada
try:
    from pdf2image import convert_from_path
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False
    print("[AVISO] pdf2image não está instalado. Prints de planilhas não serão gerados.")


# ============================================================
# CONFIGURAÇÃO BÁSICA (PATHS, .ENV, GEMINI, PASTAS)
# ============================================================

BASE_DIR = os.path.dirname(os.path.dirname(os.path.dirname(__file__)))
dotenv_path = os.path.join(BASE_DIR, ".env")
if os.path.exists(dotenv_path):
    load_dotenv(dotenv_path)

DATA_DIR = os.path.join(BASE_DIR, "data")
UPLOAD_DIR = os.path.join(DATA_DIR, "uploads")
REL_DIR = os.path.join(DATA_DIR, "relatorios")
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(REL_DIR, exist_ok=True)

# Gemini
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
GEMINI_MODEL_TEXT = os.getenv("GEMINI_MODEL_TEXT", "gemini-2.5-pro")

text_model = None
if GEMINI_API_KEY:
    print(f"[INFO] GEMINI_API_KEY detectada (prefixo={GEMINI_API_KEY[:6]}...)")
    genai.configure(api_key=GEMINI_API_KEY)
    try:
        text_model = genai.GenerativeModel(GEMINI_MODEL_TEXT)
        print(f"[INFO] Carregado modelo Gemini: {GEMINI_MODEL_TEXT}")
    except Exception as e:
        print(f"[AVISO] Falha ao carregar modelo {GEMINI_MODEL_TEXT}: {e}")
        print("[AVISO] Tentando fallback para 'gemini-1.5-pro'...")
        GEMINI_MODEL_TEXT = "gemini-1.5-pro"
        text_model = genai.GenerativeModel(GEMINI_MODEL_TEXT)
        print(f"[INFO] Fallback bem-sucedido, usando: {GEMINI_MODEL_TEXT}")
else:
    print("[AVISO] GEMINI_API_KEY não configurada. IA desativada na API.")


# ============================================================
# FASTAPI + CORS
# ============================================================

app = FastAPI(title="API Jurídica - JusReport")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # ok para MVP (depois restringe)
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Jobs em memória (Render Free reinicia, mas funciona para fluxo simples)
JOBS: Dict[str, Dict[str, Any]] = {}


# ============================================================
# MODELO P/ CORPO DO /summarize (JSON)
# ============================================================

class SummarizeRequest(BaseModel):
    question: str
    case_number: str
    action_type: str
    k: int = 50
    return_json: bool = True


# ============================================================
# HEALTH (GET + HEAD)
# ============================================================

def _health_payload() -> dict:
    env_val = os.getenv("GEMINI_API_KEY")
    return {
        "service": "api-juridica",
        "gemini_env_present": env_val is not None,
        "gemini_env_prefix": (env_val[:6] + "...") if env_val else None,
        "gemini_configured": bool(env_val),
        "gemini_model": GEMINI_MODEL_TEXT if env_val else None,
    }

@app.get("/health")
def health_get():
    return _health_payload()

@app.head("/health")
def health_head():
    # Render costuma dar HEAD /health; sem isso aparecia 405
    return JSONResponse(content=_health_payload())


# ============================================================
# INGEST / STATUS
# ============================================================

@app.post("/ingest")
async def ingest(
    files: list[UploadFile] = File(...),
    case_number: str = Form(...),
    client_id: str | None = Form(None),
):
    if not files:
        raise HTTPException(status_code=400, detail="Nenhum arquivo enviado")

    f = files[0]
    job_id = str(uuid.uuid4())

    safe_name = (f.filename or "arquivo").replace("/", "_").replace("\\", "_")
    filename = f"{job_id}__{safe_name}"
    save_path = os.path.join(UPLOAD_DIR, filename)

    content = await f.read()
    if not content:
        raise HTTPException(status_code=400, detail="Arquivo vazio")

    with open(save_path, "wb") as out:
        out.write(content)

    JOBS[job_id] = {
        "status": "done",  # ingestão simplificada
        "progress": 100,
        "detail": "Ingestão concluída",
        "file_path": save_path,
        "case_number": case_number,
        "client_id": client_id,
        "meta": {},
    }

    print(f"[INGEST] job={job_id} case={case_number} bytes={len(content)} path={save_path}")
    return {"job_id": job_id}

@app.get("/status/{job_id}")
def status(job_id: str):
    job = JOBS.get(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job não encontrado")
    return {
        "status": job["status"],
        "progress": job["progress"],
        "detail": job.get("detail", ""),
        "result": None,
    }


# ============================================================
# PDF EXTRACTION (base + hotspots) + CONTEXTO POR SEÇÃO
# ============================================================

SECTION_KEYWORDS = {
    "cabecalho": [
        "classe", "vara", "comarca", "exequente", "executado", "advogado", "oab",
        "distribuição", "distribuicao", "valor da causa", "cédula", "cedula",
        "confissão", "confissao", "contrato", "título", "titulo", "ccc", "ccb",
    ],
    "penhora": [
        "renajud", "sisbajud", "bacenjud", "infojud", "serasajud", "bloqueio",
        "penhora", "arresto", "sequestro", "constrição", "constricao", "ccs",
    ],
    "valores_planilhas": [
        "planilha", "demonstrativo", "cálculo", "calculo", "atualizado",
        "juros", "multa", "saldo devedor", "r$", "bloqueio", "sisbajud",
    ],
    "movimentacoes": [
        "decisão", "decisao", "sentença", "sentenca", "despacho", "intimação",
        "intimacao", "citação", "citacao", "juntada", "embargos", "exceção",
        "excecao", "suspensão", "suspensao", "extinção", "extincao",
    ],
    "analise_juridica": [
        "nulidade", "prescrição", "prescricao", "impugnação", "impugnacao",
        "homolog", "garantia", "hipoteca", "penhor", "aval", "fiança", "fianca",
        "terceiro", "embargos", "exceção", "excecao",
    ],
}

def _build_global_sample(full_text: str, max_chars: int) -> str:
    total_len = len(full_text)
    if total_len <= max_chars:
        return full_text

    part = max_chars // 4 or max_chars

    inicio = full_text[:part]
    mid_center = total_len // 2
    mid_start = max(0, mid_center - part // 2)
    mid_end = min(total_len, mid_start + part)
    meio = full_text[mid_start:mid_end]

    pre_final_start = max(0, total_len - (part * 2))
    pre_final_end = min(total_len, pre_final_start + part)
    pre_final = full_text[pre_final_start:pre_final_end]

    fim = full_text[-part:]

    return (
        inicio
        + "\n\n=== TRECHO CENTRAL DO PROCESSO ===\n\n"
        + meio
        + "\n\n=== TRECHO PRÉ-FINAL DO PROCESSO ===\n\n"
        + pre_final
        + "\n\n=== TRECHO FINAL DO PROCESSO ===\n\n"
        + fim
    )

def _extract_text_by_page(path: str) -> Tuple[List[str], List[Any]]:
    text_by_page: List[str] = []
    pages_obj: List[Any] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            pages_obj.append(page)
            t = page.extract_text() or ""
            text_by_page.append(t)
    return text_by_page, pages_obj

def _detect_pages_by_keywords(text_by_page: List[str], keywords: List[str], max_pages: int = 12) -> List[int]:
    hits = []
    for idx, txt in enumerate(text_by_page):
        tl = (txt or "").lower()
        if any(k in tl for k in keywords):
            hits.append(idx + 1)
            if len(hits) >= max_pages:
                break
    return hits

def _make_context_for_section(
    section_key: str,
    text_by_page: List[str],
    full_text: str,
    max_chars_section: int,
) -> Tuple[str, Dict[str, Any]]:
    """
    Monta contexto por seção:
    - pega páginas relevantes por keywords (até 12)
    - + uma amostragem global curta como “cola”
    """
    kws = SECTION_KEYWORDS.get(section_key, [])
    pages = _detect_pages_by_keywords(text_by_page, kws, max_pages=12)

    parts: List[str] = []
    for p in pages:
        raw = text_by_page[p - 1] or ""
        parts.append(f"\n\n=== PÁGINA {p} (relevante p/ {section_key}) ===\n\n{raw}")

    hotspot = "".join(parts).strip()

    # cola global curta (não 120k) — isso é o que mata 502 quando repete
    glue = _build_global_sample(full_text, max_chars=max(12000, int(max_chars_section * 0.30)))

    ctx = (
        (hotspot[: int(max_chars_section * 0.70)] if hotspot else "")
        + "\n\n=== CONTEXTO GLOBAL (AMOSTRA) ===\n\n"
        + glue
    ).strip()

    meta = {"section_pages": pages, "section_key": section_key}
    # hard cap final
    if len(ctx) > max_chars_section:
        ctx = ctx[:max_chars_section]
    return ctx, meta

def _extract_for_all_sections(path: str) -> Tuple[Dict[str, str], Dict[str, Any]]:
    """
    Mantém qualidade (MAX_PDF_CHARS=120000), MAS:
    - não manda tudo para cada agente
    - cada agente recebe um contexto menor e direcionado
    """
    text_by_page, _ = _extract_text_by_page(path)
    full_text = "\n\n".join(text_by_page)
    if not full_text.strip():
        return {}, {"error": "PDF sem texto extraível"}

    # Seu alvo de qualidade:
    env_max = int(os.getenv("MAX_PDF_CHARS", "120000"))

    # Mantém 120k como “orçamento do processo”, mas por seção usamos menos (evita 502).
    # Se quiser mais pesado, aumente para 60000 por seção (mas pode voltar 502 no Free).
    max_section = int(os.getenv("MAX_SECTION_CHARS", "45000"))

    # Mesmo que env_max seja 120k, full_text pode ser maior; não precisamos cortar aqui,
    # porque cada seção já tem seu próprio cap.
    process_meta = {
        "total_full_len": len(full_text),
        "env_max_pdf_chars": env_max,
        "max_section_chars": max_section,
    }

    contexts: Dict[str, str] = {}
    sections_meta: Dict[str, Any] = {}

    for section_key in ["cabecalho", "resumo_inicial", "penhora", "valores_planilhas", "movimentacoes", "analise_juridica"]:
        # resumo_inicial usa keywords do cabecalho (mesmo “início”) como base
        use_key = "cabecalho" if section_key == "resumo_inicial" else section_key
        ctx, meta = _make_context_for_section(use_key, text_by_page, full_text, max_chars_section=max_section)
        contexts[section_key] = ctx
        sections_meta[section_key] = meta

    return contexts, {"process_meta": process_meta, "sections_meta": sections_meta}


# ============================================================
# GEMINI CALL COM RETRY
# ============================================================

def _gemini_generate(prompt: str, retries: int = 2, backoff: float = 1.5) -> str:
    if not text_model:
        raise RuntimeError("Gemini não configurado (text_model=None).")

    last_err = None
    for attempt in range(retries + 1):
        try:
            resp = text_model.generate_content(prompt)
            return (resp.text or "").strip()
        except Exception as e:
            last_err = e
            print(f"[GEMINI] erro attempt={attempt+1}/{retries+1}: {e}")
            if attempt < retries:
                time.sleep(backoff ** attempt)
    raise last_err  # type: ignore


# ============================================================
# AGENTES (PROMPTS)
# ============================================================

def _build_tasks(case_number: str, action_type: str) -> List[Dict[str, str]]:
    return [
        {
            "key": "cabecalho",
            "title": "Cabeçalho",
            "instruction": """
Você é um assistente jurídico especialista em EXECUÇÃO DE TÍTULO EXTRAJUDICIAL.

Sua tarefa NÃO é resumir, mas sim organizar todas as informações relevantes que encontrar.

Com base EXCLUSIVAMENTE no texto abaixo, extraia o CABEÇALHO do processo, contendo, o mais completo possível:
• Número dos autos
• Classe da ação
• Vara Cível e Comarca (com número da Vara, se houver)
• Data(s) da distribuição
• Exequente (nome e CNPJ/CPF, se constar)
• Executados (nomes e CNPJ/CPF, separadamente)
• Advogados do Exequente (nome, OAB, UF – todos)
• Advogados dos Executados (nome, OAB, UF – todos)
• Valor da causa (com data se houver)
• Valor executado/atualizado na inicial (com data se houver)
• Operação/contrato (CCB, confissão etc.), número da operação, valor original, datas, garantias

Regras:
- Se não aparecer, escreva "Não informado".
- Cite fls./páginas se constarem.
- Responda em Markdown com título "## Cabeçalho" e bullets iniciando com "• ".
""",
        },
        {
            "key": "resumo_inicial",
            "title": "Resumo da Petição Inicial",
            "instruction": """
Você é um assistente jurídico especialista em EXECUÇÃO DE TÍTULO EXTRAJUDICIAL.

Faça um resumo rico em detalhes da PETIÇÃO INICIAL:
1) Partes (exequente x executados; fiadores/avalistas se houver)
2) Origem da dívida (tipo de contrato; número; datas; valor original; renegociação se houver)
3) Valores (valor executado; data; juros/multa/encargos)
4) Garantias (hipoteca/penhor/aval/fiança etc.)
5) Pedidos (em bullets)

Regras:
- Seja detalhado, sem inventar.
- Comece com "## Resumo da Petição Inicial".
""",
        },
        {
            "key": "penhora",
            "title": "Tentativas de Penhora Online e Garantias",
            "instruction": """
Você é um assistente jurídico especialista em EXECUÇÃO.

Monte "## Tentativas de Penhora Online e Garantias":
- RENAJUD / SISBAJUD(BACENJUD) / INFOJUD / SERASAJUD / outros
Para cada um: pedido? decisão? cumprimento? resultado? datas/valores/fls.

Inclua:
- Penhora de móveis / imóveis (matrícula, cartório, avaliação)
- Arresto/sequestro/indisponibilidade
- Registro de penhora em matrícula

Regra: se não encontrar, diga "Não há informação nos trechos analisados sobre ...".
""",
        },
        {
            "key": "valores_planilhas",
            "title": "Valores e Planilhas de Débito",
            "instruction": """
Você é um assistente jurídico especialista em EXECUÇÃO.

Monte "## Valores e Planilhas de Débito":
1) Valor original da operação (com datas/documentos)
2) Valor executado na inicial (com data e referência de planilha/fls.)
3) Planilhas/demonstrativos: liste cada uma (data referência, valor atualizado, fls.)
4) Tabela cronológica da evolução dos valores (mesmo se só 1 linha)
5) Bloqueios SISBAJUD/BACENJUD (data, valor, desfecho)

Regra: não invente valores/datas. Se não achar, declare.
""",
        },
        {
            "key": "movimentacoes",
            "title": "Movimentações Processuais Relevantes",
            "instruction": """
Você é um assistente jurídico especialista em EXECUÇÃO.

Monte "## Movimentações Processuais Relevantes" (linha do tempo):
• dd/mm/aaaa: ato objetivo (com fls. se houver)

Inclua pelo menos: distribuição, citação, embargos/exceções, decisões relevantes,
juntadas (planilhas/laudos), decisões sobre buscas (SISBAJUD etc.), sentenças/recursos.
""",
        },
        {
            "key": "analise_juridica",
            "title": "Análise Jurídica",
            "instruction": """
Você é um assistente jurídico especialista em EXECUÇÃO.

Monte "## Análise Jurídica" em bullets, APENAS com fatos do texto.
Se não houver dado, escreva "Não informado".

Inclua (ordem):
• Exequente
• Advogados do Exequente
• Executados
• Advogados dos Executados
• Fiadores/Avalistas/Terceiros
• Assinaturas de documentos relevantes
• Citação/Intimação
• Validade das citações
• Bens móveis em garantia
• Penhora/Arresto
• Registro da penhora
• Outras penhoras
• Planilha de cálculo (última encontrada)
• Impugnação aos cálculos
• Homologação do valor
• Avaliação de bens
• Leilão
• Manifestação de terceiros
• Exceção de pré-executividade
• Embargos à execução
• Incidentes relevantes
• Bloqueios/buscas de bens (resumo efetivo)
• Prescrição/prescrição intercorrente
• Paralisação do processo
""",
        },
    ]

def _run_agents_with_section_contexts(
    contexts: Dict[str, str],
    case_number: str,
    action_type: str,
) -> Tuple[str, Dict[str, str]]:
    tasks = _build_tasks(case_number, action_type)
    sections: Dict[str, str] = {}

    for task in tasks:
        key = task["key"]
        base_text = contexts.get(key, "") or contexts.get("cabecalho", "")

        # hard safety: se vier vazio
        if not base_text.strip():
            sections[key] = f"## {task['title']}\n\nNão informado."
            continue

        prompt = f"""{task["instruction"]}

=== TEXTO DO PROCESSO (TRECHOS SELECIONADOS) ===
\"\"\"{base_text}\"\"\""""

        print(f"[AGENTE] Rodando: {key}")
        try:
            text = _gemini_generate(prompt, retries=2, backoff=1.7)
        except Exception as e:
            # fallback: tenta reduzir o contexto e repetir 1 vez
            print(f"[AGENTE] Falhou {key}. Tentando fallback com contexto reduzido. Erro: {e}")
            small = base_text[:20000]
            prompt2 = f"""{task["instruction"]}

=== TEXTO DO PROCESSO (CONTEXTO REDUZIDO) ===
\"\"\"{small}\"\"\""""
            try:
                text = _gemini_generate(prompt2, retries=1, backoff=1.5)
            except Exception as e2:
                print(f"[AGENTE] Falhou fallback {key}: {e2}")
                text = f"## {task['title']}\n\nNão informado."

        sections[key] = (text or "").strip()

    md_parts: List[str] = [f"Sumarização da {action_type} ({case_number})\n"]
    order = ["cabecalho", "resumo_inicial", "penhora", "valores_planilhas", "movimentacoes", "analise_juridica"]
    for key in order:
        chunk = (sections.get(key) or "").strip()
        if not chunk:
            title = next(t["title"] for t in tasks if t["key"] == key)
            md_parts.append(f"## {title}\n\nNão informado.")
        else:
            md_parts.append(chunk)

    return "\n\n".join(md_parts), sections


# ============================================================
# /summarize
# ============================================================

@app.post("/summarize")
async def summarize(req: SummarizeRequest):
    try:
        case_number = req.case_number
        action_type = req.action_type

        # encontrar job
        job = None
        for j in JOBS.values():
            if j.get("case_number") == case_number:
                job = j
                break

        if not job:
            raise HTTPException(status_code=404, detail="Nenhum job encontrado para esse número de processo")

        if not GEMINI_API_KEY or not text_model:
            raise HTTPException(status_code=500, detail="Gemini não configurado na API (Render Environment)")

        file_path = job["file_path"]
        if not file_path or not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="Arquivo do job não encontrado no servidor")

        # monta contextos por seção (qualidade alta sem repetir 120k 6x)
        contexts, meta = _extract_for_all_sections(file_path)
        if not contexts:
            raise HTTPException(status_code=500, detail="Não foi possível extrair texto do PDF")

        # salva meta no job
        job_meta = job.get("meta") or {}
        job_meta.update(meta)
        job["meta"] = job_meta

        final_md, sections = _run_agents_with_section_contexts(
            contexts=contexts,
            case_number=case_number,
            action_type=action_type,
        )

        return {
            "summary_markdown": final_md,
            "sections": sections,
            "used_chunks": [],
            "result": {"meta": meta},
        }

    except HTTPException:
        raise
    except Exception as e:
        tb = traceback.format_exc()
        print("ERRO EM /summarize:\n", tb)
        raise HTTPException(status_code=500, detail=f"{e.__class__.__name__}: {e}")


# ============================================================
# /export/docx
# ============================================================

@app.post("/export/docx")
async def export_docx(
    content: str = Form(...),
    filename: str = Form("relatorio.docx"),
    case_number: str | None = Form(None),
    include_planilha_images: bool = Form(False),
):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for line in content.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        else:
            doc.add_paragraph(line)

    # prints opcionais
    if include_planilha_images and case_number and PDF2IMAGE_AVAILABLE:
        job = None
        for j in JOBS.values():
            if j.get("case_number") == case_number:
                job = j
                break

        if job:
            meta = job.get("meta") or {}
            # você pode guardar páginas de planilha aqui no futuro
            planilha_pages = (meta.get("planilha_pages") or [])
            planilha_pages = sorted(set(planilha_pages))

            file_path = job.get("file_path")
            if file_path and os.path.exists(file_path) and planilha_pages:
                try:
                    images = convert_from_path(file_path)
                    doc.add_page_break()
                    doc.add_heading("Anexos – Planilhas e Bloqueios Relevantes", level=1)

                    for p in planilha_pages:
                        if 1 <= p <= len(images):
                            img = images[p - 1]
                            img_bytes = io.BytesIO()
                            img.save(img_bytes, format="PNG")
                            img_bytes.seek(0)
                            doc.add_paragraph(f"Planilha / demonstrativo – pág. {p}")
                            doc.add_picture(img_bytes, width=Inches(6.0))
                            doc.add_paragraph("")
                except Exception as e:
                    print(f"[AVISO] Falha ao inserir imagens: {e}")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
